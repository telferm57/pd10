# -*- coding: utf-8 -*-
"""
Created on Fri Aug 18 09:12:22 2017

@author:M Telfer
"""


from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

def top5Features(series, clf): # move this to class 
    ''' formatted list of top 5 features, or 'not available'''
   
    top5_f = series.seriesStats[clf]['features']
    top5_i = series.seriesStats[clf]['importancies']
    type(top5_i)
    if isinstance(top5_i,(int)):  # -1 indicates not available 
        return 'Not Available'      
    else:
        featureImportancies = list(sorted(zip(top5_f,top5_i), key=lambda x: x[1],reverse=True))
        featureImportancies = [[x[0],"%.3f" % x[1]] for x in featureImportancies]   
        featureImportancies = featureImportancies[0:5]
        top5list =''
        for item in featureImportancies:
            top5list = top5list + ':'.join(item) + '; '
        return top5list
    
def printSeriesRegister(register):
    ''' pass in list of all series , format to a table and print in word ''' 
    document = Document()
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    document.add_heading('Series Register', 0)
    table = document.add_table(rows=1, cols=5)
    table.cell(0,0).width = 360000 # 1cm
    hdr_cells = table.rows[0].cells
    
    hdr_cells[0].text = 'Series ID'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Classes'
    hdr_cells[3].text = 'Features'
    hdr_cells[4].text = 'Classes Removed'
    
    for series, params in register.items():
     
        row_cells = table.add_row().cells
        table.cell(0,0).width = 360000 # 1cm
        row_cells[0].text = series
        row_cells[1].text = params['desc']
        row_cells[2].text = str(params['jc'])
        row_cells[3].text = str(params['excl'])
        row_cells[4].text = str(params['remcl'])
    
    document.add_page_break()
    
    document.save('.//reports//reportregister' + '.docx')
      
    return        
     
#%%    
def sampleRep(statsDict,title,fn,subtitle):
     

    from docx import Document
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENT
    
    document = Document()
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    #section.page_width = new_width
    #section.page_height = new_heigh
    document.add_heading(title, 0)
    
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    
    document.add_heading(subtitle, level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')
    
    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )
    
    #document.add_picture('monty-truth.png', width=Inches(1.25))
    # table 1 - clf 1 prformance 
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'series'
    hdr_cells[1].text = 'clf'
    hdr_cells[2].text = 'f1_macro Mean +/- 2*SD'
    hdr_cells[3].text = 'MAD'
    hdr_cells[4].text = 'Top 5 features'
    for series, statObj in statsDict.items():
        for clf in statObj.clfs:
            top5 = top5Features(statObj,clf)
            f1_m = statObj.seriesStats[clf]['f1_macro']
            f1_m_std = statObj.seriesStats[clf]['params']
            mad1 = statObj.seriesStats[clf]['MAD1']
            mad1_std = statObj.seriesStats[clf]['trainshape']
            row_cells = table.add_row().cells
            row_cells[0].text = str(statObj.series)
            row_cells[1].text = str(clf)
            row_cells[2].text = str("%.3f +/- %.3f" % (f1_m, 2*f1_m_std))
            row_cells[3].text = str("%.3f +/- %.3f" % (mad1, 2*mad1_std))
            row_cells[4].text = str(top5)
    
    document.add_page_break()
    
    document.save('.//reports//report ' + fn + '.docx')
